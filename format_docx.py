from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import re
import pandas as pd
from datetime import datetime


def parse_brief(brief_text):
    """Parse the AI-generated brief into sections using markers."""
    sections = {}
    markers = [
        "[OVERVIEW]", "[FORECAST REVIEW]", "[DATA COVERAGE]",
        "[THEMATIC ANALYSIS]", "[GEOGRAPHIC FOCUS]",
        "[TRENDS AND OUTLOOK]", "[WHAT TO WATCH]", "[REFERENCES]"
    ]
    for i, marker in enumerate(markers):
        start = brief_text.find(marker)
        if start == -1:
            continue
        start += len(marker)
        end = len(brief_text)
        for next_marker in markers[i + 1:]:
            next_pos = brief_text.find(next_marker)
            if next_pos != -1:
                end = next_pos
                break
        sections[marker] = brief_text[start:end].strip()
    return sections


def add_page_number_field(paragraph):
    """Add a PAGE field to a paragraph."""
    run = paragraph.add_run()
    fld_char_begin = parse_xml(
        '<w:fldChar {} w:fldCharType="begin"/>'.format(nsdecls("w"))
    )
    run._element.append(fld_char_begin)
    run2 = paragraph.add_run()
    instr = parse_xml(
        '<w:instrText {} xml:space="preserve"> PAGE </w:instrText>'.format(nsdecls("w"))
    )
    run2._element.append(instr)
    run3 = paragraph.add_run()
    fld_char_end = parse_xml(
        '<w:fldChar {} w:fldCharType="end"/>'.format(nsdecls("w"))
    )
    run3._element.append(fld_char_end)


def set_section_page_start(section, start_num):
    """Set the starting page number for a section."""
    sectPr = section._sectPr
    pgNumType = parse_xml(
        '<w:pgNumType {} w:start="{}"/>'.format(nsdecls("w"), start_num)
    )
    sectPr.append(pgNumType)


def add_title(doc, reporting_period):
    """Add title block."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)

    run_title = p.add_run("SOMALIA MONTHLY CONFLICT BRIEF")
    run_title.bold = True
    run_title.font.name = "Arial"
    run_title.font.size = Pt(11)

    p.add_run("\n")

    run_period = p.add_run("REPORTING PERIOD: {}".format(reporting_period.upper()))
    run_period.bold = True
    run_period.font.name = "Arial"
    run_period.font.size = Pt(11)

    p.add_run("\n")

    run_date = p.add_run("Generated: {}".format(
        datetime.now().strftime("%d %B %Y")
    ))
    run_date.font.name = "Arial"
    run_date.font.size = Pt(9)
    run_date.italic = True


def add_section_heading(doc, text):
    """Add a left-aligned, bold, all-caps section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text.upper())
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)


def add_inline_formatted_text(paragraph, text):
    """Add text to a paragraph, formatting [Comment:], [Assumption:] and **bold** inline."""
    pattern = r'(\[Comment:.*?\]|\[Assumption:.*?\]|\*\*.+?\*\*)'
    parts = re.split(pattern, text, flags=re.DOTALL)

    for part in parts:
        if not part:
            continue
        if part.startswith("[Comment:") and part.endswith("]"):
            run = paragraph.add_run(" " + part)
            run.italic = True
            run.font.name = "Arial"
            run.font.size = Pt(11)
        elif part.startswith("[Assumption:") and part.endswith("]"):
            run = paragraph.add_run(" " + part)
            run.italic = True
            run.font.name = "Arial"
            run.font.size = Pt(11)
        elif part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(11)
        else:
            run = paragraph.add_run(part)
            run.font.name = "Arial"
            run.font.size = Pt(11)


def add_body_text(doc, text):
    """Add regular body text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(text.strip())
    run.font.name = "Arial"
    run.font.size = Pt(11)


def add_sub_headed_paragraph(doc, heading, content):
    """Add a paragraph with bold sub-heading, regular full stop, then content."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    run_h = p.add_run(heading)
    run_h.bold = True
    run_h.font.name = "Arial"
    run_h.font.size = Pt(11)
    run_stop = p.add_run(".")
    run_stop.bold = False
    run_stop.font.name = "Arial"
    run_stop.font.size = Pt(11)
    if content.strip():
        add_inline_formatted_text(p, " " + content.strip())


def add_section_content(doc, content):
    """Parse section content, handling **sub-headings** and plain paragraphs."""
    paragraphs = content.split("\n\n")
    for para_text in paragraphs:
        para_text = para_text.strip()
        if not para_text:
            continue

        match = re.match(r"\*\*(.+?)\*\*\s*(.*)", para_text, re.DOTALL)
        if match:
            heading = match.group(1).strip()
            # Remove trailing full stop from heading if present
            heading = heading.rstrip(".")
            # Convert all-caps headings to sentence case
            if heading.isupper():
                heading = heading[0].upper() + heading[1:].lower()
            body = match.group(2).strip()
            body = " ".join(body.split("\n"))
            add_sub_headed_paragraph(doc, heading, body)
        else:
            text = " ".join(para_text.split("\n"))
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            add_inline_formatted_text(p, text)


def add_watch_items(doc, content):
    """Add What to Watch items as separate paragraphs with bold heading and regular full stop."""
    lines = content.strip().split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # Remove leading numbers, dashes or bullets
        line = re.sub(r"^[\d]+[\.\)]\s*", "", line)
        line = re.sub(r"^[-\*\u2022]\s*", "", line)
        if not line:
            continue

        # Try to split on first full stop or colon to create a bold heading
        # Look for a pattern like "Bold heading: rest of text" or "Bold heading. Rest of text"
        heading_match = re.match(r"^(.+?)[:\.](.+)$", line)
        if heading_match and len(heading_match.group(1)) < 80:
            heading = heading_match.group(1).strip()
            # Remove **markers** if present
            heading = re.sub(r"\*\*(.+?)\*\*", r"\1", heading)
            body = heading_match.group(2).strip()
            add_sub_headed_paragraph(doc, heading, body)
        else:
            # No clear heading structure, just add as body text
            line = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(12)
            add_inline_formatted_text(p, line)


def add_references(doc, content):
    """Add the references section with smaller font."""
    lines = content.strip().split("\n")
    for line in lines:
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(line)
        run.font.name = "Arial"
        run.font.size = Pt(9)


def add_annex_section(doc, annex_letter, annex_title):
    """Add a new section for an annex with proper formatting."""
    new_section = doc.add_section()
    new_section.top_margin = Cm(2)
    new_section.bottom_margin = Cm(2)
    new_section.left_margin = Cm(2)
    new_section.right_margin = Cm(2)

    set_section_page_start(new_section, 1)

    footer = new_section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prefix_run = fp.add_run("{}-".format(annex_letter))
    prefix_run.font.name = "Arial"
    prefix_run.font.size = Pt(11)
    add_page_number_field(fp)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("ANNEX {}: {}".format(annex_letter, annex_title.upper()))
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(11)


def build_methodology_note(doc):
    """Build methodology and limitations annex."""
    methodology = [
        (
            "Conflict data",
            "Armed Conflict Location and Event Data (ACLED), a disaggregated "
            "conflict data collection project that records political violence, "
            "demonstrations and select non-violent political "
            "developments globally. ACLED data is updated weekly and accessed "
            "via API. Event-level data for the reporting month is provided to "
            "the analytical model; baseline months are provided as aggregate "
            "statistics only."
        ),
        (
            "Food security data",
            "IPC (Integrated Phase Classification) phase assessments, accessed "
            "via the IPC API. Provides population-level food security phase "
            "classifications by admin1 region. IPC phases: 1 = Minimal, "
            "2 = Stressed, 3 = Crisis, 4 = Emergency, 5 = Famine. IPC data is "
            "treated as supplementary context only and not used as a primary "
            "analytical driver."
        ),
        (
            "Rainfall data",
            "CHIRPS (Climate Hazards Group InfraRed Precipitation with Station "
            "data) dekadal rainfall anomaly data, downloaded from the "
            "Humanitarian Data Exchange (HDX). Provides the 3-month rainfall "
            "anomaly (r3q) as a percentage of the 1989-2018 long-term average "
            "by admin1 region. Values below 80% indicate drought conditions; "
            "above 130% indicates heavy rainfall. Data may be labelled "
            "provisional until confirmed as final."
        ),
        (
            "Population data",
            "UNFPA 2021 population projections, aggregated to admin1. Used to "
            "calculate per-capita conflict event and fatality rates for the "
            "reporting month. Per-capita rates are expressed as events or "
            "fatalities per 100,000 population. These are projections, not "
            "census data, and should be treated as indicative."
        ),
        (
            "Displacement data",
            "Primary source: UNHCR/OCHA Harmonised IDP Figures, downloaded "
            "from the Humanitarian Data Exchange (HDX). This dataset "
            "consolidates IDP figures from IOM DTM and CCCM cluster sources "
            "at district level and is aggregated to admin1 for this product. "
            "Regions absent from the harmonised file are supplemented with "
            "IOM DTM V3 API data. Figures should be treated as minimum "
            "estimates as coverage varies by region."
        ),
        (
            "Analytical process",
            "The brief is produced using a structured analytical pipeline. "
            "ACLED event data for the reporting period and baseline months "
            "is pulled via API. The reporting month's event-level data is "
            "provided to a large language model (Claude, Anthropic) alongside "
            "IPC, CHIRPS rainfall, UNFPA population and IOM DTM displacement "
            "datasets, plus context notes on terminology and actor definitions. "
            "The model derives its own analytical conclusions from the data, "
            "following a detailed prompt that specifies structure, style, "
            "analytical discipline and Somalia-specific checks."
        ),
        (
            "AI generation",
            "The narrative analysis is generated by an AI model. It is not "
            "human-authored. The model derives its own analytical conclusions "
            "from the event data, constrained to what the data can support. "
            "It is instructed to distinguish factual statements from analytical "
            "commentary using [Comment: ...] blocks, present competing "
            "assumptions for significant claims, use calibrated probability "
            "language, and cite ACLED event IDs as footnote references. "
            "Non-ACLED data sources are attributed in-text but not footnoted."
        ),
        (
            "Probability language",
            "The brief uses calibrated probability terms for forward-looking "
            "judgments. Highly likely: 75-90% probability. Likely: 55-75%. "
            "Roughly even: 45-55%. Unlikely: 25-45%. Highly unlikely: 10-25%. "
            "The model calibrates conservatively and defaults one tier lower "
            "than its initial assessment. Tactical predictions (military and "
            "security patterns) may use the full scale. Political predictions "
            "default to 'roughly even' because political outcomes depend on "
            "factors not captured in event data."
        ),
        (
            "Limitations",
            "This tool has significant limitations. It has no access to human "
            "source reporting, political intelligence or direct observation. "
            "ACLED data is subject to reporting bias: events in accessible "
            "urban areas are more likely to be captured than events in remote "
            "or armed-group-controlled areas. Fatality figures are often "
            "estimates from single sources. Rainfall and displacement data may "
            "lag real conditions on the ground. Population projections are from "
            "2021 and may not reflect current distribution. The AI model may "
            "produce plausible-sounding analysis not fully supported by the "
            "data. Forward-looking judgments are preliminary. All assessments "
            "require human review."
        ),
        (
            "Recommended use",
            "This brief is a situational awareness tool, not a substitute for "
            "human analytical judgment. It provides a structured first draft "
            "that a human analyst can review, edit and build upon. It should "
            "not be used as the sole basis for operational decisions."
        ),
    ]
    for heading, content in methodology:
        add_sub_headed_paragraph(doc, heading, content)


def create_brief_docx(brief_text, df, reporting_period="March 2026",
                       output_path="somalia_brief.docx"):
    """Create a formatted Word document from the brief and data."""
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(12)
    style.paragraph_format.line_spacing = 1.15

    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number_field(fp)

    add_title(doc, reporting_period)

    sections = parse_brief(brief_text)

    section_map = [
        ("[OVERVIEW]", "OVERVIEW"),
        ("[FORECAST REVIEW]", "FORECAST REVIEW"),
        ("[DATA COVERAGE]", "DATA COVERAGE"),
        ("[THEMATIC ANALYSIS]", "THEMATIC ANALYSIS"),
        ("[GEOGRAPHIC FOCUS]", "GEOGRAPHIC FOCUS"),
        ("[TRENDS AND OUTLOOK]", "TRENDS AND OUTLOOK"),
    ]

    for marker, heading in section_map:
        if marker not in sections:
            continue
        add_section_heading(doc, heading)
        add_section_content(doc, sections[marker])

    if "[WHAT TO WATCH]" in sections:
        add_section_heading(doc, "WHAT TO WATCH")
        add_watch_items(doc, sections["[WHAT TO WATCH]"])

    if "[REFERENCES]" in sections:
        add_section_heading(doc, "REFERENCES")
        add_references(doc, sections["[REFERENCES]"])

    # Annex A: Methodology only (data summary removed)
    add_annex_section(doc, "A", "Methodology and Limitations")
    build_methodology_note(doc)

    doc.save(output_path)
    print("Word document saved: {}".format(output_path))
    return output_path


if __name__ == "__main__":
    print("format_docx.py loaded successfully.")
